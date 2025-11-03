# Professional Word document certificate generator
import sys
sys.path.insert(0, '/usr/lib/python3/dist-packages')
sys.path.insert(0, '/home/cj/.local/lib/python3.12/site-packages')

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from datetime import datetime
    import os

    def generate_docx_certificate(student_name, course_title, cert_type="course", student_username=None, completed_modules=None):
        """Generate a professional Word document certificate"""
        doc = Document()

        # Set up styles
        title_style = doc.styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = RGBColor(0, 123, 255)  # Blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle_style = doc.styles.add_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Pt(16)
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name_style = doc.styles.add_style('NameStyle', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.size = Pt(24)
        name_style.font.color.rgb = RGBColor(40, 167, 69)  # Green
        name_style.font.bold = True
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        normal_style = doc.styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_paragraph("🎓 CERTIFICATE OF COMPLETION 🎓", style='TitleStyle')
        title.paragraph_format.space_after = Pt(30)

        # Platform name
        platform = doc.add_paragraph("SKILLSWAP LEARNING PLATFORM", style='SubtitleStyle')
        platform.paragraph_format.space_after = Pt(20)

        # Certification text
        doc.add_paragraph("This is to certify that", style='NormalStyle')

        # Student name
        name_para = doc.add_paragraph(student_name.upper(), style='NameStyle')
        name_para.paragraph_format.space_after = Pt(10)

        # Username
        if student_username:
            username_para = doc.add_paragraph(f"Username: {student_username}", style='NormalStyle')
            username_para.paragraph_format.space_after = Pt(15)

        # Course completion text
        doc.add_paragraph("has successfully completed the comprehensive course:", style='NormalStyle')

        # Course title
        course_title_para = doc.add_paragraph(f'"{course_title.upper()}"', style='NormalStyle')
        course_title_para.runs[0].bold = True
        course_title_para.runs[0].font.color.rgb = RGBColor(0, 123, 255)  # Blue
        course_title_para.paragraph_format.space_after = Pt(20)

        # Course details
        details = [
            f"Course Type: {cert_type.title()}",
            f"Modules Completed: {completed_modules or 'All Required Modules'}",
            "Final Assessment: PASSED",
            f"Completion Date: {datetime.now().strftime('%B %d, %Y')}",
            f"Certificate ID: SS-{datetime.now().strftime('%Y%m%d')}-{hash(student_name + course_title) % 10000:04d}"
        ]

        for detail in details:
            doc.add_paragraph(detail, style='NormalStyle')

        # Validation text
        validation_para = doc.add_paragraph()
        validation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        validation_para.add_run("""
This certificate validates that the recipient has demonstrated proficiency in the subject matter
through comprehensive learning, practical exercises, and successful completion of all required assessments.
""")

        # Signature section
        doc.add_paragraph("")  # Empty line
        signature_line = doc.add_paragraph("___________________________________", style='NormalStyle')
        signature_title = doc.add_paragraph("Authorized Signature", style='NormalStyle')
        platform_name = doc.add_paragraph("SkillSwap Learning Platform", style='NormalStyle')
        date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'), style='NormalStyle')

        return doc

    def save_certificate(student_name, course_title, file_path, cert_type="course", student_username=None, completed_modules=None):
        """Save Word document certificate to file system"""
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # Default values if not provided
        if student_username is None:
            student_username = student_name.lower().replace(' ', '_')
        if completed_modules is None:
            completed_modules = "All Required Modules"

        doc = generate_docx_certificate(student_name, course_title, cert_type, student_username, completed_modules)
        doc.save(file_path)
        return file_path

except ImportError as e:
    print(f"python-docx not available: {e}, falling back to text certificates")
    # Fallback to text-based certificate
    from datetime import datetime
    import os

    def save_certificate(student_name, course_title, file_path, cert_type="course", student_username=None, completed_modules=None):
        """Generate a simple text-based certificate"""
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # Default values if not provided
        if student_username is None:
            student_username = student_name.lower().replace(' ', '_')
        if completed_modules is None:
            completed_modules = "All Required Modules"

        certificate_content = f"""
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║                          CERTIFICATE OF COMPLETION                           ║
║                                                                              ║
║                                                                              ║
║              This certifies that                                             ║
║                                                                              ║
║                                                                              ║
║                          {student_name.upper()}                              ║
║                                                                              ║
║                          Username: {student_username}                        ║
║                                                                              ║
║                                                                              ║
║              has successfully completed the                                  ║
║                                                                              ║
║                                                                              ║
║                          {course_title}                                      ║
║                                                                              ║
║                          Course Type: {cert_type.title()}                    ║
║                          Modules Completed: {completed_modules}              ║
║                          Final Assessment: PASSED                             ║
║                          Completion Date: {datetime.now().strftime('%B %d, %Y')} ║
║                                                                              ║
║                          SkillSwap Learning Platform                         ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(certificate_content)

        return file_path